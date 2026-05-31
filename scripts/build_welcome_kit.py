from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("artifacts/welcome-kit/SrS_Logics_Client_Welcome_Kit.docx")

NAVY = RGBColor(9, 17, 29)
NAVY_SOFT = RGBColor(19, 28, 44)
NAVY_PANEL = "111B2B"
NAVY_PANEL_2 = "16253C"
CYAN = RGBColor(166, 240, 255)
BLUE = RGBColor(31, 146, 255)
VIOLET = RGBColor(106, 97, 255)
INK = RGBColor(33, 47, 68)
TEXT = RGBColor(74, 91, 116)
WHITE = RGBColor(255, 255, 255)
PALE = RGBColor(223, 234, 250)
LINE = "D9E5F5"
SOFT = "F6FAFF"
SOFT_2 = "EEF4FF"
SOFT_3 = "F2F0FF"


def set_page_geometry(section) -> None:
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.34)
    section.footer_distance = Inches(0.28)


def set_run_font(run, *, name="Arial", size=11, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def format_paragraph(paragraph, *, before=0, after=0, line=1.15, align=None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_table_borders(table, color=LINE, size=8) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_page_number(paragraph) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fld.append(OxmlElement("w:r"))
    paragraph._p.append(fld)


def configure_header_footer(section) -> None:
    header = section.header.paragraphs[0]
    format_paragraph(header, after=0)
    a = header.add_run("SRS LOGICS")
    set_run_font(a, size=9.5, color=NAVY, bold=True)
    b = header.add_run("  |  Client Welcome Kit")
    set_run_font(b, size=9.5, color=TEXT)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    format_paragraph(footer, before=0, after=0)
    a = footer.add_run("Page ")
    set_run_font(a, size=9, color=TEXT)
    add_page_number(footer)


def add_cover(doc: Document) -> None:
    hero = doc.add_table(rows=2, cols=2)
    hero.alignment = WD_TABLE_ALIGNMENT.LEFT
    hero.autofit = False
    hero.columns[0].width = Inches(4.45)
    hero.columns[1].width = Inches(2.1)
    remove_table_borders(hero)

    top_left = hero.cell(0, 0)
    top_right = hero.cell(0, 1)
    bottom_left = hero.cell(1, 0)
    bottom_right = hero.cell(1, 1)

    for cell, fill in [
        (top_left, "0B1220"),
        (top_right, NAVY_PANEL),
        (bottom_left, NAVY_PANEL_2),
        (bottom_right, "1A2140"),
    ]:
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=170, start=180, bottom=170, end=170)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p = top_left.paragraphs[0]
    format_paragraph(p, after=8)
    r = p.add_run("SRS LOGICS")
    set_run_font(r, size=10.5, color=CYAN, bold=True)

    p = top_left.add_paragraph()
    format_paragraph(p, after=10, line=1.0)
    r = p.add_run("Client\nWelcome Kit")
    set_run_font(r, size=24, color=WHITE, bold=True)

    p = top_left.add_paragraph()
    format_paragraph(p, after=0, line=1.28)
    r = p.add_run(
        "A modern onboarding packet for custom software engagements, approvals, testing, deployment, and handover."
    )
    set_run_font(r, size=11.2, color=PALE)

    p = top_right.paragraphs[0]
    format_paragraph(p, after=8)
    r = p.add_run("ENGAGEMENT FILE")
    set_run_font(r, size=9.2, color=CYAN, bold=True)

    meta = [
        ("Prepared for", "[Client Name / Company Name]"),
        ("Project", "[Software / System Name]"),
        ("Prepared by", "SrS Logics"),
        ("Issued on", date.today().strftime("%d %B %Y")),
    ]
    for label, value in meta:
        p = top_right.add_paragraph()
        format_paragraph(p, after=8, line=1.16)
        a = p.add_run(f"{label}\n")
        set_run_font(a, size=8.5, color=RGBColor(154, 176, 214), bold=True)
        b = p.add_run(value)
        set_run_font(b, size=10.1, color=WHITE, bold=True)

    p = bottom_left.paragraphs[0]
    format_paragraph(p, after=6)
    r = p.add_run("WORKING PRINCIPLE")
    set_run_font(r, size=9.2, color=CYAN, bold=True)
    p = bottom_left.add_paragraph()
    format_paragraph(p, after=0, line=1.2)
    r = p.add_run(
        "The software should reflect the way the business actually runs. This kit is built to reduce ambiguity and create a disciplined start."
    )
    set_run_font(r, size=10.4, color=PALE)

    p = bottom_right.paragraphs[0]
    format_paragraph(p, after=6)
    r = p.add_run("DELIVERY WINDOW")
    set_run_font(r, size=9.2, color=CYAN, bold=True)
    p = bottom_right.add_paragraph()
    format_paragraph(p, after=0, line=1.2)
    r = p.add_run("Most engagements move from requirement gathering to testing within roughly 4 to 6 weeks.")
    set_run_font(r, size=10.4, color=PALE)


def add_signal_strip(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(2.15)
    table.columns[1].width = Inches(2.15)
    table.columns[2].width = Inches(2.15)
    remove_table_borders(table)

    data = [
        ("Requirement clarity", "Real business process understanding before development starts."),
        ("Testing discipline", "Structured review, milestone checks, and sign-off handling."),
        ("Deployment readiness", "Clean handover with support continuity after launch."),
    ]
    fills = ["EAF5FF", "EFF0FF", "EAF5FF"]
    accents = [BLUE, VIOLET, BLUE]
    for idx, (title, text) in enumerate(data):
        cell = table.cell(0, idx)
        set_cell_shading(cell, fills[idx])
        set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
        p = cell.paragraphs[0]
        format_paragraph(p, after=4)
        a = p.add_run(f"{title}\n")
        set_run_font(a, size=10, color=accents[idx], bold=True)
        b = p.add_run(text)
        set_run_font(b, size=10.1, color=INK)


def add_section_heading(doc: Document, kicker: str, title: str) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, before=18, after=3)
    a = p.add_run(kicker.upper())
    set_run_font(a, size=9.3, color=BLUE, bold=True)
    p = doc.add_paragraph()
    format_paragraph(p, after=8, line=1.0)
    b = p.add_run(title)
    set_run_font(b, size=18, color=NAVY, bold=True)


def add_body(doc: Document, text: str, *, after=6) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, after=after, line=1.24)
    r = p.add_run(text)
    set_run_font(r, size=10.9, color=INK)


def add_info_cards(doc: Document, cards: list[tuple[str, str]], fills: list[str] | None = None) -> None:
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(3.15)
    table.columns[1].width = Inches(3.15)
    remove_table_borders(table)
    fills = fills or [SOFT, SOFT_3, SOFT_2, SOFT]
    for idx, (title, copy) in enumerate(cards):
        cell = table.cell(idx // 2, idx % 2)
        set_cell_shading(cell, fills[idx])
        set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
        p = cell.paragraphs[0]
        format_paragraph(p, after=4)
        a = p.add_run(f"{title}\n")
        set_run_font(a, size=10.3, color=VIOLET if idx % 2 else BLUE, bold=True)
        b = p.add_run(copy)
        set_run_font(b, size=10.3, color=INK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_two_column_panels(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(3.15)
    table.columns[1].width = Inches(3.15)
    remove_table_borders(table)

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    for cell, fill in [(left, SOFT_2), (right, SOFT_3)]:
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=150, start=170, bottom=150, end=170)

    p = left.paragraphs[0]
    format_paragraph(p, after=5)
    r = p.add_run("How we work")
    set_run_font(r, size=10.5, color=BLUE, bold=True)
    bullets = [
        "Delivery stays tied to the approved business workflow, not generic software assumptions.",
        "Client response time affects delivery speed, especially during requirement confirmation and testing closure.",
        "The best software outcomes happen when process owners stay involved in reviews.",
    ]
    for item in bullets:
        p = left.add_paragraph(style="List Bullet")
        format_paragraph(p, after=3, line=1.14)
        r = p.add_run(item)
        set_run_font(r, size=10.3, color=INK)

    p = right.paragraphs[0]
    format_paragraph(p, after=5)
    r = p.add_run("Commercial and support")
    set_run_font(r, size=10.5, color=VIOLET, bold=True)
    bullets = [
        "Payments are generally structured in three phases: confirmation, API/testing milestone, and handover.",
        "Infrastructure choice and performance scope affect the final commercial size of the build.",
        "Post-handover support is typically included for 6 months unless revised in writing.",
    ]
    for item in bullets:
        p = right.add_paragraph(style="List Bullet")
        format_paragraph(p, after=3, line=1.14)
        r = p.add_run(item)
        set_run_font(r, size=10.3, color=INK)


def add_timeline_table(doc: Document) -> None:
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(0.9)
    table.columns[1].width = Inches(1.55)
    table.columns[2].width = Inches(2.35)
    table.columns[3].width = Inches(1.75)
    set_table_borders(table)
    headers = ["Stage", "Typical timing", "What happens", "Client role"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "EAF5FF" if i % 2 == 0 else "EFF0FF")
        set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
        p = cell.paragraphs[0]
        format_paragraph(p, after=0)
        r = p.add_run(text)
        set_run_font(r, size=9.8, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])

    rows = [
        ("01", "Week 1", "Requirement understanding, operating flow review, scope alignment, and approval of the working plan.", "Share workflow, decision makers, and priority outcomes."),
        ("02", "Week 2", "Interface structure, process definition, database planning, and internal build setup.", "Confirm process details, fields, and exception cases."),
        ("03", "Weeks 3-4", "Core software development, logic implementation, reporting flow, and platform assembly.", "Review previews and respond to clarifications quickly."),
        ("04", "Weeks 4-6", "Testing, API validation, issue fixes, deployment preparation, and structured handover.", "Participate in testing and prepare final users for go-live."),
    ]
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            set_cell_margins(cell, top=95, start=130, bottom=95, end=130)
            p = cell.paragraphs[0]
            format_paragraph(p, after=0, line=1.14)
            r = p.add_run(text)
            set_run_font(r, size=10.1, color=INK, bold=(col_idx == 0))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_checklist_table(doc: Document) -> None:
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(3.4)
    table.columns[1].width = Inches(1.45)
    table.columns[2].width = Inches(1.55)
    set_table_borders(table)
    headers = ["Required from client side", "Suggested owner", "Status"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "EAF5FF" if i % 2 == 0 else "EFF0FF")
        set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
        p = cell.paragraphs[0]
        format_paragraph(p, after=0)
        r = p.add_run(text)
        set_run_font(r, size=9.8, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Confirmed project contact and approval authority", "Management", "[ ]"),
        ("Current workflow notes or live operating examples", "Operations", "[ ]"),
        ("Reporting, dashboard, or analysis expectations", "Management / finance", "[ ]"),
        ("Existing data samples, formats, or exports", "Client technical/data", "[ ]"),
        ("Testing users and final sign-off participants", "Operations / management", "[ ]"),
    ]
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            set_cell_margins(cell, top=95, start=130, bottom=95, end=130)
            p = cell.paragraphs[0]
            format_paragraph(
                p,
                after=0,
                line=1.14,
                align=WD_ALIGN_PARAGRAPH.CENTER if col_idx == 2 else None,
            )
            r = p.add_run(text)
            set_run_font(r, size=10.1, color=INK)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_handover_panel(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.45)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "0F1C30")
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    p = cell.paragraphs[0]
    format_paragraph(p, after=6)
    r = p.add_run("Typical handover package")
    set_run_font(r, size=11.2, color=CYAN, bold=True)

    items = [
        "Approved software build aligned to the agreed scope",
        "Deployment-ready environment and relevant database structure",
        "Role handling, access logic, and operating instructions where required",
        "Testing closure notes and final issue summary",
        "Defined support window after handover",
    ]
    for item in items:
        p = cell.add_paragraph(style="List Bullet")
        format_paragraph(p, after=3, line=1.14)
        r = p.add_run(item)
        set_run_font(r, size=10.4, color=PALE)


def build_document(output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    set_page_geometry(section)
    configure_header_footer(section)

    add_cover(doc)
    add_signal_strip(doc)

    add_section_heading(doc, "Welcome", "A Clear Start To The Engagement")
    add_body(
        doc,
        "Thank you for choosing SrS Logics. This welcome kit is designed to make the software engagement feel structured from the beginning. It aligns both teams on how the project will move, what approvals are needed, what should be ready from the client side, and what the handover journey looks like before development starts.",
    )

    add_section_heading(doc, "Snapshot", "Engagement At A Glance")
    add_info_cards(
        doc,
        [
            ("Engagement model", "Founder-led delivery with direct business understanding and software shaped around your actual operating process."),
            ("Timeline window", "Most software engagements move from requirement gathering to testing within roughly 4 to 6 weeks."),
            ("Support window", "Post-handover support is typically included for 6 months unless revised in the final commercial note."),
            ("Commercial structure", "Payments are generally handled in three phases: confirmation, testing milestone, and final handover."),
        ],
    )

    add_section_heading(doc, "Operating Model", "How The Delivery Will Work")
    add_two_column_panels(doc)

    add_section_heading(doc, "Timeline", "Typical Delivery Journey")
    add_timeline_table(doc)

    add_section_heading(doc, "Kickoff", "What We Need From Client Side")
    add_body(
        doc,
        "Projects move faster when ownership is clear, operating examples are available early, and approvals are not fragmented across multiple channels. The checklist below is the recommended starting pack for kickoff.",
        after=8,
    )
    add_checklist_table(doc)

    add_section_heading(doc, "Handover", "What Closure Typically Includes")
    add_handover_panel(doc)

    add_section_heading(doc, "Final Note", "Shared Responsibility Creates Better Software")
    add_body(
        doc,
        "The purpose of this kit is not only to introduce the engagement, but to create a better working rhythm. When scope, approvals, testing, and handover are handled clearly from the start, the software becomes more stable, more useful, and more aligned with the way the business actually operates.",
        after=3,
    )
    add_body(
        doc,
        "Use this document as the working orientation pack for the kickoff discussion, and update the placeholders once the client-specific project details are confirmed.",
        after=0,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    build_document(OUTPUT_PATH)
